from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from pypdf import PdfReader

MAX_RESUME_BYTES = 5 * 1024 * 1024
MAX_RESUME_PAGES = 20
MAX_RESUME_TEXT_CHARS = 50_000
MAX_DOCX_UNCOMPRESSED_BYTES = 20 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class ResumeParseError(ValueError):
    pass


def _pdf_text(data: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(data))
        if reader.is_encrypted:
            raise ResumeParseError("Encrypted PDFs are not supported")
        if len(reader.pages) > MAX_RESUME_PAGES:
            raise ResumeParseError(f"Resume must be {MAX_RESUME_PAGES} pages or fewer")
        parts: list[str] = []
        length = 0
        for page in reader.pages:
            page_text = page.extract_text() or ""
            parts.append(page_text[: MAX_RESUME_TEXT_CHARS - length])
            length += len(page_text)
            if length >= MAX_RESUME_TEXT_CHARS:
                break
        return "\n".join(parts)
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError("Could not read that PDF") from exc


def _docx_text(data: bytes) -> str:
    try:
        with ZipFile(BytesIO(data)) as archive:
            if sum(item.file_size for item in archive.infolist()) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ResumeParseError("The DOCX expands beyond the safe processing limit")
        document = Document(BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError("Could not read that DOCX file") from exc


def extract_resume_text(data: bytes, filename: str) -> str:
    if not data:
        raise ResumeParseError("The uploaded file is empty")
    if len(data) > MAX_RESUME_BYTES:
        raise ResumeParseError("Resume must be 5 MB or smaller")

    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ResumeParseError("Upload a PDF or DOCX resume")
    if extension == ".pdf" and not data.startswith(b"%PDF"):
        raise ResumeParseError("The file does not appear to be a PDF")
    if extension == ".docx" and not data.startswith(b"PK"):
        raise ResumeParseError("The file does not appear to be a DOCX file")

    text = _pdf_text(data) if extension == ".pdf" else _docx_text(data)
    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not normalized:
        raise ResumeParseError("No readable text was found in the resume")
    return normalized[:MAX_RESUME_TEXT_CHARS]
